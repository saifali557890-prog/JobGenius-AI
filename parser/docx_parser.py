import docx
import io
from utils.logger import get_logger

logger = get_logger("DOCXParser")

class DOCXParser:
    @staticmethod
    def extract_text(file_bytes_or_path) -> str:
        """
        Extracts raw text from Microsoft Word (.docx) files including headers and tables.
        """
        try:
            if isinstance(file_bytes_or_path, (bytes, bytearray)):
                doc = docx.Document(io.BytesIO(file_bytes_or_path))
            else:
                doc = docx.Document(file_bytes_or_path)

            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract text embedded inside tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
                        
            logger.info("DOCX text extraction completed successfully.")
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}")
            return ""